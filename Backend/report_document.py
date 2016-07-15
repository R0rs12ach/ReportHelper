#encoding=utf-8
# -*-coding:gb2312-*-
"""
@AUTHOR: R0rs12ach@gmail.com
@Time: 2016-07-14 22:00
@Deployment: 当前脚本的运行依赖于docx，而docx依赖于lxml
@Description: 如果你不想使用数据库自动抽取本周工作内容进行发送，也可将当前文件当做一个脚本使用
"""
from docx import Document
from docx.shared import Inches
import time

# 创建一个doc文档
document = Document()

# 定义文档题目
document.add_heading(u'xxx部门工作周报', 0)

# 当你没有数据库的时候，你可以直接在这里填写周报内容, 有数据库的情况下可以读数据库内容
recordset = [{"id":'1',"content":u'工作内容条目一',"level":u'紧急重要',"status":u'未完成'},{"id":'2',"content":u'工作内容条目二',"level":u'重要不紧急',"status":u'完成'},{"id":'3',"content":u'工作内容条目三',"level":u'紧急重要',"status":u'完成'},{"id":"4","content":u'工作内容条目四','level':u'紧急重要','status':u'完成'}]

nextset = [{"id":'1',"content":u'计划工作内容条目一',"level":u'紧急且重要',"status":u'尚未开始'}]

# 周报的日期和周报的汇报人
basic_info = u"时间：" + time.strftime("%Y-%m-%d",time.localtime()) + u"            姓名：R0rs12ach"
p = document.add_paragraph(basic_info)

# 本周工作情况总结
this_week_title = document.add_paragraph('', style='ListBullet')
this_week_title.add_run(u'本周工作完成情况').bold = True

# 设定过去的一周的表格样式为蓝色主题
table_this_week = document.add_table(rows=1, cols=4, style='MediumShading1-Accent1')
this_hdr_cells = table_this_week.rows[0].cells
this_hdr_cells[0].text = u'序号'
this_hdr_cells[1].text = u'工作内容'
this_hdr_cells[2].text = u'重要/紧急程度'
this_hdr_cells[3].text = u'完成情况'
for item in recordset:
    this_row_cells = table_this_week.add_row().cells
    this_row_cells[0].text = str(item["id"])
    this_row_cells[1].text = item["content"]
    this_row_cells[2].text = item["level"]
    this_row_cells[3].text = item["status"]

# 加一个分隔，免得两个表间隔太近，不太美观，人，一定要看美观整洁的，否则，跟咸鱼有什么区别
document.add_paragraph()

# 下周工作情况展望
next_week_title = document.add_paragraph('', style='ListBullet')
next_week_title.add_run(u'下周工作计划').bold = True

# 设定即将到来的一周的表格样式为绿色主题
table_next_week = document.add_table(rows=1, cols=4, style='MediumShading1-Accent3')
next_hdr_cells = table_next_week.rows[0].cells
next_hdr_cells[0].text = u'序号'
next_hdr_cells[1].text = u'工作内容'
next_hdr_cells[2].text = u'重要/紧急程度'
next_hdr_cells[3].text = u'完成情况'
for item in nextset:
    next_row_cells = table_next_week.add_row().cells
    next_row_cells[0].text = str(item["id"])
    next_row_cells[1].text = item["content"]
    next_row_cells[2].text = item["level"]
    next_row_cells[3].text = item["status"]

# 如果你希望分页，可以直接执行下面函数
# document.add_page_break()

#可配置为你想要保存的word文档文件名
document.save(u'R0rs12ach的工作周报.docx')