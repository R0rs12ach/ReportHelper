#encoding=utf-8
# -*-coding:gb2312-*-
from docx import Document
from docx.shared import Inches
import time

# 创建一个doc文档
document = Document()

# 定义文档题目
document.add_heading(u'UED部门工作周报', 0)

recordset = [{"id":'1',"content":u'开发一个自动周报系统，智能更新内容，定时推送',"level":u'非紧急但重要',"status":u'完成'},{"id":'2',"content":u'给UED组进行技术培训，为后续网站的技术支持培养新人',"level":u'紧急且重要',"status":u'完成'},{"id":'3',"content":u'研究FRP开发模式',"level":u'非紧急但重要',"status":u'完成'},{"id":"4","content":u'配合测试维护CONE及其插件','level':u'紧急重要','status':u'完成'}]

nextset = [{"id":'1',"content":u'在windows Server 环境上开发vmware的cone插件部分',"level":u'紧急且重要',"status":u'尚未开始'}]

# 周报的日期和周报的汇报人
basic_info = u"时间：" + time.strftime("%Y-%m-%d",time.localtime()) + u"            姓名：彭岁万"
p = document.add_paragraph(basic_info)

# 本周工作情况总结
this_week_title = document.add_paragraph('', style='ListBullet')
this_week_title.add_run(u'本周工作完成情况').bold = True

# 设定过去的一周的表格样式为蓝色主题
table_this_week = document.add_table(rows=1, cols=4, style='MediumShading1-Accent1')
this_hdr_cells = table_this_week.rows[0].cells
this_hdr_cells[0].text = u'序号'
this_hdr_cells[1].text = u'工作内容'
this_hdr_cells[2].text = u'重要/紧急程度'
this_hdr_cells[3].text = u'完成情况'
for item in recordset:
    this_row_cells = table_this_week.add_row().cells
    this_row_cells[0].text = str(item["id"])
    this_row_cells[1].text = item["content"]
    this_row_cells[2].text = item["level"]
    this_row_cells[3].text = item["status"]

# 加一个分隔，免得两个表间隔太近，不太美观，人，一定要看美观整洁的，否则，跟咸鱼有什么区别
document.add_paragraph()

# 下周工作情况展望
next_week_title = document.add_paragraph('', style='ListBullet')
next_week_title.add_run(u'下周工作计划').bold = True

# 设定即将到来的一周的表格样式为绿色主题
table_next_week = document.add_table(rows=1, cols=4, style='MediumShading1-Accent3')
next_hdr_cells = table_next_week.rows[0].cells
next_hdr_cells[0].text = u'序号'
next_hdr_cells[1].text = u'工作内容'
next_hdr_cells[2].text = u'重要/紧急程度'
next_hdr_cells[3].text = u'完成情况'
for item in nextset:
    next_row_cells = table_next_week.add_row().cells
    next_row_cells[0].text = str(item["id"])
    next_row_cells[1].text = item["content"]
    next_row_cells[2].text = item["level"]
    next_row_cells[3].text = item["status"]


# document.add_page_break()

document.save(u'彭岁万的工作周报.docx')
